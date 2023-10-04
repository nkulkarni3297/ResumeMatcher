import pandas as pd
import streamlit as st
import openai
from difflib import SequenceMatcher
from docx import Document
from io import BytesIO
import fitz  # PyMuPDF
import base64


# Set your OpenAI API key here
with st.sidebar:
    openai.api_key = st.text_input('Add your OpenAI API Key', type='password')
data = []

# Function to calculate similarity ratio between strings
def similarity_ratio(a, b):
    return SequenceMatcher(None, a, b).ratio()

def get_completion(prompt, model="gpt-3.5-turbo-16k"):
    messages = [{"role": "user", "content": prompt}]
    
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0.7,
    )
    
    return response.choices[0].message["content"]

st.title("Resume Matcher")

job_description= st.text_area("Enter Job Description")
# File upload for job description
job_description_file = st.file_uploader("Upload Job Description", type=["txt", "pdf", "docx"])

# File upload for resume
resume_files = st.file_uploader("Upload Your Resume", type=["txt", "pdf", "docx"],accept_multiple_files=True)

if st.button("Generate Report"):
    # Check if no files are selected
    if len(resume_files) == 0:
        st.write("Please Select File")
    
    # Read content from the uploaded job description file
    
    
    # Check if it's a DOCX file
    if job_description_file is not None:
            job_description_content = job_description_file.read()
    if job_description_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(BytesIO(job_description_content))
                job_description = "\n".join(para.text for para in doc.paragraphs)
    elif job_description_file.type == "application/pdf":
                pdf = fitz.open(stream=job_description_content, filetype="pdf")
                job_description = ""
     for page_num in range(pdf.page_count):
                    page = pdf[page_num]
                    job_description += page.get_text()
     else:
                # Assume it's a text file
                job_description = job_description_content.decode("utf-8")


    # for resume_file in resume_files: 
    i=0
    while i < len(resume_files):
    # Read content from the uploaded resume file
        resume_file = resume_files[i]
        resume_content = resume_file.read()

    # Check if it's a DOCX file
    if resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(BytesIO(resume_content))
        resume_input = "\n".join(para.text for para in doc.paragraphs)
    elif resume_file.type == "application/pdf":
        pdf = fitz.open(stream=resume_content, filetype="pdf")
        resume_input = ""
        for page_num in range(pdf.page_count):
            page = pdf[page_num]
            resume_input += page.get_text()
    else:
        # Assume it's a text file
        resume_input = resume_content.decode("utf-8")

    input_text = f""" Your task is to compare the resume with the job_description provided. Provide the results in terms of the percentage of suitability for the job. And provide me the reason on what basis you have provided the percentage.
            job_description: {job_description}
            resume: {resume_input}
        """
    generated_resume = get_completion(input_text)
    data.append({"Name": resume_file.name, "Match percentage": generated_resume})
    i = i + 1

if len(resume_files) == 0:
    st.write("Please Select File")
else:
    df = pd.DataFrame(data)
    st.write(df)

    csv = df.to_csv(index=False)
    b64 = base64.b64encode(csv.encode()).decode()  # Encode to base64
    href = f'<a href="data:file/csv;base64,{b64}" download="data.csv"> Click On For Download CSV File</a>'
    st.markdown(href, unsafe_allow_html=True)


